import re
import requests
from io import BytesIO
from pdfminer.high_level import extract_text
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdftypes import resolve1
from docx import Document


def get_file_id(drive_url: str) -> str:
    patterns = [
        r"/file/d/([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)"
    ]

    for pattern in patterns:
        match = re.search(pattern, drive_url)
        if match:
            return match.group(1)

    raise ValueError("Could not extract Google Drive file ID")


def download_drive_file(drive_url: str) -> bytes:
    file_id = get_file_id(drive_url)

    download_url = (
        f"https://drive.google.com/uc?export=download&id={file_id}"
    )

    response = requests.get(download_url)
    response.raise_for_status()

    return response.content


def extract_hyperlinks(content_bytes: bytes) -> list:
    links = []
    try:
        parser = PDFParser(BytesIO(content_bytes))
        doc = PDFDocument(parser)
        for page in PDFPage.create_pages(doc):
            if 'Annots' in page.attrs:
                annots = resolve1(page.attrs['Annots'])
                if annots:
                    # Sometimes annots is a single object, sometimes a list
                    if not isinstance(annots, list):
                        annots = [annots]
                    for annot in annots:
                        annot_obj = resolve1(annot)
                        if isinstance(annot_obj, dict) and 'A' in annot_obj:
                            action = resolve1(annot_obj['A'])
                            if isinstance(action, dict) and 'URI' in action:
                                uri = action['URI']
                                if isinstance(uri, bytes):
                                    links.append(uri.decode('utf-8', 'ignore'))
    except Exception as e:
        print(f"Failed to extract links: {e}")
    return list(set(links))


def extract_resume_text(drive_url: str, file_type: str):
    content = download_drive_file(drive_url)

    if file_type.lower() == "pdf":
        text = extract_text(BytesIO(content))
        links = extract_hyperlinks(content)
        if links:
            text += "\n--- EXTRACTED HYPERLINKS ---"
            for link in links:
                text += f"- {link}\n"
        return text

    elif file_type.lower() == "docx":
        doc = Document(BytesIO(content))
        return "\n".join(
            para.text for para in doc.paragraphs
        )

    elif file_type.lower() == "txt":
        return content.decode("utf-8", errors="ignore")

    else:
        raise ValueError(
            f"Unsupported file type: {file_type}"
        )

def drive_file_content(link):
  text = extract_resume_text(link, "pdf")
  print()
  return text
